from docx import Document
from os import getcwd, makedirs, listdir, devnull
from os.path import join, exists, dirname
from shutil import copy, rmtree, make_archive, move
from lxml import etree
from zipfile import ZipFile
from docx2pdf import convert as docx_to_pdf
from pdf2image import convert_from_path
from threading import Thread
import logging
import sys

poppler_path = join(getcwd(), 'poppler_bin', 'Release', 'poppler', 'Library', 'bin')

# Generate All Certificate
def generate_certificate(reg_name: list, filename: str, output_type: list, keyValue_pairs: list):

    print(f'\n. . . GENERATING CERTIFICATES . . .\n')
    print(f'Document Files - ', end = '')

    document_path = join(getcwd(), 'custom_template', filename) 

    publicFolder = join(getcwd(), 'public')
    makedirs(publicFolder, exist_ok=True)

    docxFolder = join(publicFolder, 'docx')
    pdfFolder = join(publicFolder, 'pdf')
    imgFolder = join(publicFolder, 'img')

    makedirs(docxFolder, exist_ok=True)
    makedirs(pdfFolder, exist_ok=True)
    makedirs(imgFolder, exist_ok=True)

    try:

        for i, name in enumerate(reg_name):

            # Make a copy
            saveto_path = join(
                docxFolder, 
                f'{'_'.join([name[0].split()[-1], name[0].split()[0]]).upper()}.docx'
            )
            
            copy(document_path, saveto_path)

            # Replace Keywords
            fileCopy = Document(saveto_path)

            # body
            for paragraph in fileCopy.paragraphs:
                full_text = ''.join(run.text for run in paragraph.runs)
                for j, kv in enumerate(keyValue_pairs):
                    # Pre-determine Values
                    if j == 0:
                        for k, key in enumerate(kv):
                            if key in full_text and key:
                                full_text = full_text.replace(key, reg_name[i][k])
                                for run in paragraph.runs:
                                    run.text = ''
                                paragraph.runs[0].text = full_text
                        continue

                    if kv[0] in full_text:
                        full_text = full_text.replace(kv[0], kv[1])
                        for run in paragraph.runs:
                            run.text = ''
                        paragraph.runs[0].text = full_text


            # table
            for table in fileCopy.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for j, kv in enumerate(keyValue_pairs):
                            if j == 0:
                                for k, key in enumerate(kv):
                                    if key in cell.text and key:
                                        cell.text = cell.text.replace(key, reg_name[i][k])
                                continue
                            if kv in cell.text:
                                cell.text = cell.text.replace(key, name[0])

            fileCopy.save(saveto_path)
            

            # inline
            inlineReplaceOne(saveto_path, saveto_path, keyValue_pairs, reg_name[i])
        
        print(f'Sucessful\n')
        print(f'PDF Files:\n')

        convert_documents(docxFolder, pdfFolder, imgFolder, output_type)

    except Exception as e:
        print(f'Error saving file: {e}')
        return
    

def convert_documents(docxFolder, pdfFolder, imgFolder, output_type):
    docx_files = [f for f in listdir(docxFolder) if f.endswith('.docx')]

    def convert_to_pdf(docx_file):
        docx_path = join(docxFolder, docx_file)
        pdf_path = join(pdfFolder, docx_file.replace('.docx', '.pdf'))
        docx_to_pdf(docx_path, pdf_path)

    def convert_to_img(pdf_file):
        pdf_path = join(pdfFolder, pdf_file)
        images = convert_from_path(pdf_path, poppler_path=poppler_path)
        img_path = join(imgFolder, pdf_file.replace('.pdf', '.png'))
        images[0].save(img_path, 'PNG')

    threads = []

    if "PDF" in output_type or "IMG" in output_type:

        for docx_file in docx_files:
            thread = Thread(target=convert_to_pdf, args=(docx_file,))
            threads.append(thread)
            thread.start()

        for thread in threads:
            thread.join()
        
        print(f'\nSuccessful\n')

    if "IMG" in output_type:
        print(f'PNG Images - ', end = '')

        pdf_files = [f for f in listdir(pdfFolder) if f.endswith('.pdf')]
        threads = []
        for pdf_file in pdf_files:
            thread = Thread(target=convert_to_img, args=(pdf_file,))
            threads.append(thread)
            thread.start()

        for thread in threads:
            thread.join()

        print(f'Successful\n')

    if "IMG" in output_type and "DOCX" not in output_type and "PDF" not in output_type:
        # Delete DOCX and PDF files if only images are needed
        rmtree(docxFolder)
        rmtree(pdfFolder)

    


""" - - - - - - - - Single Certificate - - - - - - - - """

def generate_one_certificate(reg_name: list, filename: str, keyValue_pairs: list = [['[name]', '[honor]', '[quarter]']]):

    publicFolder = join(getcwd(), 'public')
    makedirs(publicFolder, exist_ok=True)

    docxFolder = join(publicFolder, 'docx')
    pdfFolder = join(publicFolder, 'pdf')
    imgFolder = join(publicFolder, 'img')

    makedirs(docxFolder, exist_ok=True)
    makedirs(pdfFolder, exist_ok=True)
    makedirs(imgFolder, exist_ok=True)

    document_path = join(getcwd(), 'custom_template', filename) 

    temporaryFolder = join(getcwd(), 'temporary')
    makedirs(temporaryFolder, exist_ok=True)

    try:

        # Make a copy
        saveto_path = join(temporaryFolder, 'preview_img.docx')
        copy(document_path, saveto_path)

        # Replace Keywords
        fileCopy = Document(saveto_path)
        
        # body
        for paragraph in fileCopy.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            for i, kv in enumerate(keyValue_pairs):
                if i == 0:
                    for j, key in enumerate(kv):
                        if key in full_text and key:
                            full_text = full_text.replace(key, reg_name[0][j])
                            for run in paragraph.runs:
                                run.text = ''
                            paragraph.runs[0].text = full_text
                    continue

                if kv[0] in full_text:
                    full_text = full_text.replace(kv[0], kv[1])
                    for run in paragraph.runs:
                        run.text = ''
                    paragraph.runs[0].text = full_text

        # table
        for table in fileCopy.tables:
            for row in table.rows:
                for cell in row.cells:
                    for i, kv in enumerate(keyValue_pairs):
                        if i == 0:
                            for j, key in enumerate(kv): 
                                if key in cell.text and key:
                                    cell.text = cell.text.replace(key, reg_name[0][j])
                                    continue

                        if kv[0] in cell.text:
                            cell.text = cell.text.replace(kv[0], kv[0])
    
        fileCopy.save(saveto_path)

        # inline
        inlineReplaceOne(saveto_path, saveto_path, keyValue_pairs, reg_name[0])

        return saveto_path

    except Exception as e:
        print(f'Error saving file [generate_certificate_script]: {e}')
        return


def inlineReplaceOne(doc_path, save_path, keywords, person):
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Ensure directories for extraction exist
    extract_path = "extracted"
    if exists(extract_path):
        rmtree(extract_path)
    makedirs(extract_path)

    # Extract the DOCX file
    with ZipFile(doc_path, 'r') as zip_ref:
        zip_ref.extractall(extract_path)

    # Parse and modify document.xml
    document_xml_path = join(extract_path, 'word', 'document.xml')
    with open(document_xml_path, 'r', encoding='utf-8') as file:
        tree = etree.parse(file)

    # Find text boxes in the document
    for text_box in tree.xpath('.//w:drawing//w:txbxContent', namespaces=namespaces):
        paragraphs = text_box.findall('.//w:p', namespaces=namespaces)
        for paragraph in paragraphs:
            runs = paragraph.findall('.//w:r', namespaces=namespaces)
            full_text = ''.join(run.find('.//w:t', namespaces=namespaces).text or '' for run in runs)

            # Replace keywords in text
            for i, kv in enumerate(keywords):
                if i == 0:
                    for j, key in enumerate(kv):
                        if key in full_text and key:
                            full_text = full_text.replace(key, person[j])
                            continue

                if kv[0] in full_text:
                    full_text = full_text.replace(kv[0], kv[1])

            # Update the text in runs
            for run in runs:
                text_element = run.find('.//w:t', namespaces=namespaces)
                if text_element is not None:
                    text_element.text = ''  # Clear old text
            if runs and len(runs) > 0 and runs[0].find('.//w:t', namespaces=namespaces) is not None:
                runs[0].find('.//w:t', namespaces=namespaces).text = full_text

    # Write modified XML back to document.xml
    with open(document_xml_path, 'wb') as file:
        tree.write(file, xml_declaration=True, encoding='utf-8', pretty_print=True)

    # Repackage the modified contents back into a DOCX file
    make_archive(save_path.replace('.docx', ''), 'zip', extract_path)
    move(save_path.replace('.docx', '.zip'), save_path)

    # Clean up extracted files
    rmtree(extract_path)









    



